#!/usr/bin/env python3
"""从 PDF/docx/txt 简历抽取纯文本。用法: python parse_resume.py <路径>
依赖(按需): pip install pdfplumber python-docx --break-system-packages"""
import sys, os
def pdf(p):
    import pdfplumber
    with pdfplumber.open(p) as f: return "\n".join((pg.extract_text() or "") for pg in f.pages).strip()
def docx_(p):
    import docx; d=docx.Document(p)
    parts=[x.text for x in d.paragraphs]
    for t in d.tables:
        for r in t.rows: parts.append("\t".join(c.text for c in r.cells))
    return "\n".join(parts).strip()
def txt(p):
    return open(p,encoding="utf-8",errors="replace").read().strip()
def main():
    if len(sys.argv)!=2: sys.exit("用法: python parse_resume.py <简历路径>")
    p=sys.argv[1]; e=os.path.splitext(p)[1].lower()
    if not os.path.isfile(p): sys.exit("文件不存在: "+p)
    t = pdf(p) if e==".pdf" else docx_(p) if e in (".docx",".doc") else txt(p) if e in (".txt",".md") else sys.exit("不支持: "+e)
    print(t or sys.exit("空文本(可能扫描件),请改用视觉读图或粘贴文本"))
if __name__=="__main__": main()
